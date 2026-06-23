"""Minimal Markdown -> .docx converter for the paper draft.

Handles the subset of Markdown used in planning_failover.md: ATX headings,
fenced code blocks (rendered monospace), blockquotes (italic notes), ordered and
unordered lists, horizontal rules, and inline **bold** / `code`.

Usage (ephemeral python-docx, no venv pollution):
    uv run --with python-docx python docs/paper/md_to_docx.py \
        docs/paper/planning_failover.md docs/paper/planning_failover.docx
"""

from __future__ import annotations

import re
import sys
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt

_INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def _add_inline(paragraph, text: str) -> None:
    for tok in _INLINE.split(text):
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Courier New"
        elif tok:
            paragraph.add_run(tok)


def convert(src: str, dst: str) -> None:
    doc = Document()
    in_code = False
    code: list[str] = []

    for raw in open(src, encoding="utf-8").read().splitlines():
        line = raw.rstrip()

        if line.lstrip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code))
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                code = []
            in_code = not in_code
            continue
        if in_code:
            code.append(line)
            continue

        if not line.strip() or line.strip() == "---":
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.add_run(line[2:].strip()).italic = True
        elif re.match(r"^\d+\.\s", line):
            _add_inline(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s", "", line))
        elif line.startswith("- "):
            _add_inline(doc.add_paragraph(style="List Bullet"), line[2:])
        else:
            _add_inline(doc.add_paragraph(), line)

    # Stamp the document's own created/modified dates (the python-docx template
    # otherwise leaves a 2013 default that Word displays as the document date).
    now = datetime.now(timezone.utc)
    doc.core_properties.created = now
    doc.core_properties.modified = now

    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
